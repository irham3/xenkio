from docx import Document # type: ignore
from docx.shared import Pt, Inches, Emu # type: ignore
from docx.oxml import OxmlElement # type: ignore
from docx.oxml.ns import qn # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH # type: ignore
from docx.enum.table import WD_TABLE_ALIGNMENT # type: ignore
import pdfplumber # type: ignore
import pypdf # type: ignore
import io
from PIL import Image # type: ignore

# --- Constants ---
PDF_PT_TO_WORD_PT = 1.0       # 1 PDF point = 1 Word point
PDF_PT_TO_INCHES = 1.0 / 72.0 # 1 inch = 72 PDF points

# --- Helpers ---

def rgb_to_hex(rgb_tuple):
    if not rgb_tuple or len(rgb_tuple) < 3:
        return 'FFFFFF'
    r, g, b = rgb_tuple[0], rgb_tuple[1], rgb_tuple[2]
    if isinstance(r, float) and r <= 1.0:
        r, g, b = int(r * 255), int(g * 255), int(b * 255)
    return f"{r:02x}{g:02x}{b:02x}".upper()

def set_cell_background(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell_valign(cell, align='center'):
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data is not None:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_cell_width(cell, width_inches):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')

def set_row_height(table_obj, row_idx, height_pt):
    tr = table_obj.rows[row_idx]._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    trHeight = trPr.find(qn('w:trHeight'))
    if trHeight is None:
        trHeight = OxmlElement('w:trHeight')
        trPr.append(trHeight)
    trHeight.set(qn('w:val'), str(int(height_pt * 20)))
    trHeight.set(qn('w:hRule'), 'atLeast')

def disable_table_autofit(table_obj):
    tbl = table_obj._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    autoFit = tblPr.find(qn('w:tblLayout'))
    if autoFit is None:
        autoFit = OxmlElement('w:tblLayout')
        tblPr.append(autoFit)
    autoFit.set(qn('w:type'), 'fixed')

def set_paragraph_spacing(p, before=0, after=0):
    pPr = p._p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p._p.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'), str(int(before * 20)))
    spacing.set(qn('w:after'), str(int(after * 20)))

def set_cell_margins(cell, top=0, bottom=0, left=1, right=1):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement('w:{}'.format(edge))
        el.set(qn('w:w'), str(int(val * 20)))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def overlap_area(bbox1, bbox2):
    x_left = max(bbox1[0], bbox2[0])
    y_top = max(bbox1[1], bbox2[1])
    x_right = min(bbox1[2], bbox2[2])
    y_bottom = min(bbox1[3], bbox2[3])
    if x_right < x_left or y_bottom < y_top:
        return 0.0
    return (x_right - x_left) * (y_bottom - y_top)

def bbox_area(bbox):
    return max(0, bbox[2] - bbox[0]) * max(0, bbox[3] - bbox[1])

def map_font(font_name):
    if not font_name: return 'Arial'
    
    # Strip PDF font prefixes (e.g. ABCDEF+)
    if '+' in font_name:
        font_name = font_name.split('+')[-1]
    
    base = font_name.split(',')[0].split('-')[0]
    
    # User requested exact font, so we use the exact name
    # We still clean up styling suffixes
    base = base.replace('Bold', '').replace('Italic', '').replace('Oblique', '')
    
    if 'Times' in base:
        return 'Times New Roman'
    return base

def is_bold(font_name):
    if not font_name: return False
    fl = font_name.lower()
    return 'bold' in fl or 'heavy' in fl or 'black' in fl

def is_italic(font_name):
    if not font_name: return False
    fl = font_name.lower()
    return 'italic' in fl or 'oblique' in fl

def is_line_dashed(line):
    dash = line.get('dash')
    if dash is None:
        return False
    if isinstance(dash, (list, tuple)):
        if len(dash) >= 1:
            d = dash[0] if not isinstance(dash[0], (list, tuple)) else dash[0]
            if isinstance(d, (list, tuple)) and len(d) >= 2:
                return True
            elif isinstance(dash, tuple) and len(dash) == 2 and isinstance(dash[0], list):
                return len(dash[0]) >= 2
    return False

def classify_border(cell_bbox, edge, all_lines):
    """Classify each edge of a cell as 'solid', 'dotted', or 'nil'."""
    is_horizontal = edge in ('top', 'bottom')
    
    if edge == 'top':    coord = cell_bbox[1]
    elif edge == 'bottom': coord = cell_bbox[3]
    elif edge == 'left':   coord = cell_bbox[0]
    else:                  coord = cell_bbox[2]
    
    matching_lines = []
    for l in all_lines:
        if is_horizontal:
            if abs(l['top'] - coord) < 4 and l['x0'] < cell_bbox[2] and l['x1'] > cell_bbox[0]:
                matching_lines.append(l)
        else:
            if abs(l['x0'] - coord) < 4 and l['top'] < cell_bbox[3] and l['bottom'] > cell_bbox[1]:
                matching_lines.append(l)
    
    if not matching_lines:
        return 'nil'
    
    for l in matching_lines:
        if is_line_dashed(l):
            return 'dotted'
    
    return 'single'

def apply_run_style(run, word):
    run.font.size = Pt(round(word['size']))
    run.font.name = map_font(word.get('fontname'))
    if is_bold(word.get('fontname')):
        run.bold = True
    if is_italic(word.get('fontname')):
        run.italic = True

def detect_cell_alignments(cell_bbox, cell_words):
    """Detect horizontal and vertical alignment of text within a cell."""
    if not cell_words:
        return 'left', 'top'
    
    cx0, cy0, cx1, cy1 = cell_bbox
    cell_w = cx1 - cx0
    cell_h = cy1 - cy0
    
    text_x0 = min(w['x0'] for w in cell_words)
    text_x1 = max(w['x1'] for w in cell_words)
    text_y0 = min(w['top'] for w in cell_words)
    text_y1 = max(w['bottom'] for w in cell_words)
    
    left_margin = text_x0 - cx0
    right_margin = cx1 - text_x1
    top_margin = text_y0 - cy0
    bottom_margin = cy1 - text_y1
    
    # Horizontal
    if cell_w > 0 and abs(left_margin - right_margin) < max(10, cell_w * 0.15):
        h_align = 'center'
    else:
        h_align = 'left'
    
    # Vertical
    if cell_h > 0 and abs(top_margin - bottom_margin) < max(5, cell_h * 0.15):
        v_align = 'center'
    elif top_margin < bottom_margin:
        v_align = 'top'
    else:
        v_align = 'bottom'
    
    return h_align, v_align

def add_words_to_paragraph(p, line_words, cell_bbox=None, h_align=None):
    if not line_words:
        return
    
    if h_align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif h_align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # else: default LEFT
    
    last_x1 = line_words[0]['x0']
    for w in line_words:
        gap = w['x0'] - last_x1
        if gap > 8:
            # Add a custom tab stop exactly where the word should start
            # tab stops are relative to the paragraph's left edge (margin + indent)
            # if we are inside a cell, margin is 0. If free text, margin is 36.
            # It's safer to just let docx handle it by adding a tab_stop
            # We'll calculate relative to line_words[0]['x0'] for simplicity if it's the same paragraph
            rel_pos = w['x0'] - line_words[0]['x0']
            if rel_pos > 0:
                try:
                    p.paragraph_format.tab_stops.add_tab_stop(Pt(rel_pos))
                except Exception:
                    pass
            tab_run = p.add_run('\t')
            tab_run.font.size = Pt(round(w['size']))
        
        run = p.add_run(w['text'] + ' ')
        apply_run_style(run, w)
        last_x1 = w['x1']

def get_words_in_bbox(words, bbox, threshold=0.5):
    result = []
    for w in words:
        w_bbox = (w['x0'], w['top'], w['x1'], w['bottom'])
        w_area = bbox_area(w_bbox)
        if w_area > 0 and overlap_area(bbox, w_bbox) > threshold * w_area:
            result.append(w)
    return result

def group_words_into_lines(cell_words):
    if not cell_words:
        return []
    lines = {}
    for w in cell_words:
        y_key = round(w['top'] / 3) * 3
        if y_key not in lines:
            lines[y_key] = []
        lines[y_key].append(w)
    result = []
    for y in sorted(lines.keys()):
        line = sorted(lines[y], key=lambda w: w['x0'])
        result.append(line)
    return result


# --- Main Converter ---

def convert_pdf_to_word(input_path, output_path):
    doc = Document()
    
    # Configure document to use A4 and exact margins based on PDF mediabox
    for section in doc.sections:
        section.page_width = Pt(595)
        section.page_height = Pt(842)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    word_count = 0
    reader = pypdf.PdfReader(input_path)
    
    with pdfplumber.open(input_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            pypdf_page = reader.pages[page_num]
            tables = page.find_tables()
            table_bboxes = [t.bbox for t in tables]
            words = page.extract_words(extra_attrs=['fontname', 'size'])
            all_rects = page.rects
            all_lines = page.lines
            all_pdf_lines = all_lines  # lines used for border analysis
            
            # Collect valid images from pypdf (skip grayscale masks)
            valid_images = []
            plumber_images = page.images
            p_idx = 0
            for img in pypdf_page.images:
                try:
                    im = Image.open(io.BytesIO(img.data))
                    if im.mode == 'L' or 'Mask' in img.name:
                        continue
                    bbox = None
                    if p_idx < len(plumber_images):
                        pi = plumber_images[p_idx]
                        bbox = (pi['x0'], pi['top'], pi['x1'], pi['bottom'])
                    p_idx += 1
                    valid_images.append({'data': img.data, 'bbox': bbox})
                except Exception:
                    pass
            
            # Build ordered elements
            elements = []
            
            dashed_lines = [l for l in all_lines if is_line_dashed(l)]
            if dashed_lines:
                # Group dashed lines by proximity to form separate bounding boxes
                clusters = []
                for l in dashed_lines:
                    b1 = (l['x0'], l['top'], l['x1'], l['bottom'])
                    matched = []
                    for i, c in enumerate(clusters):
                        # check intersection with 5pt margin
                        if any(not (b1[2]+5 < b2['x0'] or b1[0]-5 > b2['x1'] or b1[3]+5 < b2['top'] or b1[1]-5 > b2['bottom']) for b2 in c):
                            matched.append(i)
                    if not matched:
                        clusters.append([l])
                    else:
                        clusters[matched[0]].append(l)
                        for i in reversed(matched[1:]):
                            clusters[matched[0]].extend(clusters[i])
                            del clusters[i]
                
                for c in clusters:
                    d_x0 = min(l['x0'] for l in c)
                    d_y0 = min(l['top'] for l in c)
                    d_x1 = max(l['x1'] for l in c)
                    d_y1 = max(l['bottom'] for l in c)
                    dashed_bbox = (d_x0, d_y0, d_x1, d_y1)
                    
                    # Check if this area is already covered by a pdfplumber table
                    already_in_table = False
                    for tb in table_bboxes:
                        if overlap_area(dashed_bbox, tb) > 0.5 * bbox_area(dashed_bbox):
                            already_in_table = True
                            break
                    
                    if not already_in_table:
                        # Find unique X and Y coordinates to build a grid
                        x_coords = sorted(set(round(l['x0'], 0) for l in c if abs(l['top'] - l['bottom']) > 5)
                                         | set(round(l['x1'], 0) for l in c if abs(l['top'] - l['bottom']) > 5)
                                         | {round(d_x0, 0), round(d_x1, 0)})
                        y_coords = sorted(set(round(l['top'], 0) for l in c if abs(l['x1'] - l['x0']) > 5)
                                         | set(round(l['bottom'], 0) for l in c if abs(l['x1'] - l['x0']) > 5)
                                         | {round(d_y0, 0), round(d_y1, 0)})
                        
                        if len(x_coords) >= 2 and len(y_coords) >= 2:
                            elements.append({
                                'type': 'dashed_table',
                                'top': d_y0,
                                'bbox': dashed_bbox,
                                'x_coords': x_coords,
                                'y_coords': y_coords,
                            })
                            table_bboxes.append(dashed_bbox)
            
            # Detect large standalone rectangles (like "Catatan" box) missed by find_tables
            tables_to_process = list(tables)
            for r in page.rects:
                w = r['width']
                h = r['height']
                if w > 300 and h > 50:
                    rect_bbox = (r['x0'], r['top'], r['x1'], r['bottom'])
                    
                    overlap = False
                    for tb in table_bboxes:
                        if overlap_area(rect_bbox, tb) > 0.5 * bbox_area(rect_bbox):
                            overlap = True
                            break
                    for e in elements: # check dashed_table overlaps
                        if e['type'] == 'dashed_table' and overlap_area(rect_bbox, e['bbox']) > 0.5 * bbox_area(rect_bbox):
                            overlap = True
                            break
                            
                    if not overlap:
                        class FakeRow:
                            def __init__(self, bbox):
                                self.cells = [bbox]
                        class FakeTable:
                            def __init__(self, bbox):
                                self.bbox = bbox
                                self.rows = [FakeRow(bbox)]
                        
                        tables_to_process.append(FakeTable(rect_bbox))
                        table_bboxes.append(rect_bbox)
            
            # Tables
            for t in tables_to_process:
                elements.append({
                    'type': 'table',
                    'data': t,
                    'top': t.bbox[1],
                    'bbox': t.bbox
                })
            
            # Free text (not in any table)
            text_lines_map = {}
            for w in words:
                word_count += 1
                in_table = False
                for tb in table_bboxes:
                    if tb[0] - 2 <= w['x0'] and tb[1] - 2 <= w['top'] and tb[2] + 2 >= w['x1'] and tb[3] + 2 >= w['bottom']:
                        in_table = True
                        break
                if not in_table:
                    y = round(w['top'] / 3) * 3
                    if y not in text_lines_map:
                        text_lines_map[y] = []
                    text_lines_map[y].append(w)
            
            for y in sorted(text_lines_map.keys()):
                line_words = sorted(text_lines_map[y], key=lambda w: w['x0'])
                # Calculate bbox for line
                x0 = line_words[0]['x0']
                x1 = line_words[-1]['x1']
                y0 = min(w['top'] for w in line_words)
                y1 = max(w['bottom'] for w in line_words)
                elements.append({
                    'type': 'text',
                    'words': line_words,
                    'top': y0,
                    'bbox': (x0, y0, x1, y1)
                })
            
            elements.sort(key=lambda e: e['top'])
            
            # ===== Render elements =====
            image_placed = set()
            last_bottom = 0
            
            for e in elements:
                space_before = max(0, e['top'] - last_bottom) if last_bottom > 0 else 0
                
                # --- DASHED TABLE (Nomor Dokumen / Revisi / Halaman) ---
                if e['type'] == 'dashed_table':
                    x_coords = e['x_coords']
                    y_coords = e['y_coords']
                    n_rows = len(y_coords) - 1
                    n_cols = len(x_coords) - 1
                    
                    if n_rows <= 0 or n_cols <= 0:
                        continue
                    
                    table_obj = doc.add_table(rows=n_rows, cols=n_cols)
                    table_obj.style = 'Table Grid'
                    table_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
                    disable_table_autofit(table_obj)
                    
                    for ci in range(n_cols):
                        col_w = (x_coords[ci + 1] - x_coords[ci]) * PDF_PT_TO_INCHES
                        table_obj.columns[ci].width = Inches(col_w)
                    
                    for ri in range(n_rows):
                        row_h = y_coords[ri + 1] - y_coords[ri]
                        set_row_height(table_obj, ri, row_h)
                        for ci in range(n_cols):
                            cell_bbox = (x_coords[ci], y_coords[ri], x_coords[ci + 1], y_coords[ri + 1])
                            cell = table_obj.cell(ri, ci)
                            set_cell_width(cell, (x_coords[ci + 1] - x_coords[ci]) * PDF_PT_TO_INCHES)
                            
                            # Detect actual alignment
                            d_cw = get_words_in_bbox(words, cell_bbox)
                            d_h, d_v = detect_cell_alignments(cell_bbox, d_cw)
                            set_cell_valign(cell, d_v)
                            
                            # All borders are dotted, but top border is nil (no top border in original)
                            border_kw = {}
                            for edge_name in ('top', 'bottom', 'left', 'right'):
                                cls = classify_border(cell_bbox, edge_name, all_pdf_lines)
                                if cls == 'dotted':
                                    border_kw[edge_name] = {"val": "dotted", "sz": "4", "color": "000000", "space": "0"}
                                elif cls == 'nil':
                                    border_kw[edge_name] = {"val": "nil"}
                                else:
                                    border_kw[edge_name] = {"val": "single", "sz": "4", "color": "000000", "space": "0"}
                            set_cell_border(cell, **border_kw)
                            
                            cell.text = ''
                            if d_cw:
                                lines = group_words_into_lines(d_cw)
                                for li, line in enumerate(lines):
                                    p = cell.paragraphs[0] if li == 0 else cell.add_paragraph()
                                    add_words_to_paragraph(p, line, cell_bbox, h_align=d_h)
                    
                    doc.add_paragraph().paragraph_format.space_after = Pt(2)
                
                # --- REGULAR TABLE ---
                elif e['type'] == 'table':
                    t = e['data']
                    rows = len(t.rows)
                    cols = max((len(row.cells) for row in t.rows), default=0)
                    if cols == 0:
                        continue
                    
                    table_obj = doc.add_table(rows=rows, cols=cols)
                    table_obj.style = 'Table Grid'
                    table_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
                    disable_table_autofit(table_obj)
                    
                    # Calculate unique column X positions from the first row with all cells
                    col_xs = []
                    for row in t.rows:
                        xs = []
                        for c in row.cells:
                            if c is not None:
                                xs.append(c[0])
                        if len(xs) == cols:
                            col_xs = xs
                            break
                    
                    # Set column widths from bbox
                    for j in range(cols):
                        # Find any row that has this column defined
                        for row in t.rows:
                            if row.cells[j] is not None:
                                cw = (row.cells[j][2] - row.cells[j][0]) * PDF_PT_TO_INCHES
                                table_obj.columns[j].width = Inches(cw)
                                break
                    
                    # Set row heights
                    for i, row in enumerate(t.rows):
                        for c in row.cells:
                            if c is not None:
                                rh = c[3] - c[1]
                                set_row_height(table_obj, i, rh)
                                break
                    
                    # Merge cells using coordinate-based detection
                    # This correctly handles both vertical (rowspan) and horizontal (colspan) merges
                    
                    # Build midpoint coordinates for each column and row
                    col_mid_x = [None] * cols
                    for j in range(cols):
                        for i in range(rows):
                            c = t.rows[i].cells[j]
                            if c is not None:
                                col_mid_x[j] = (c[0] + c[2]) / 2.0
                                break
                    row_mid_y = [None] * rows
                    for i in range(rows):
                        for j in range(cols):
                            c = t.rows[i].cells[j]
                            if c is not None:
                                row_mid_y[i] = (c[1] + c[3]) / 2.0
                                break
                    
                    # For each None cell, find which non-None cell's bbox contains its midpoint
                    master_grid = [[None]*cols for _ in range(rows)]
                    for i in range(rows):
                        for j in range(cols):
                            if t.rows[i].cells[j] is not None:
                                master_grid[i][j] = (i, j)
                            else:
                                px = col_mid_x[j]
                                py = row_mid_y[i]
                                if px is None or py is None:
                                    # Fallback: look up then left
                                    if i > 0 and master_grid[i-1][j] is not None:
                                        master_grid[i][j] = master_grid[i-1][j]
                                    elif j > 0 and master_grid[i][j-1] is not None:
                                        master_grid[i][j] = master_grid[i][j-1]
                                    else:
                                        master_grid[i][j] = (i, j)
                                    continue
                                
                                found = False
                                for mi in range(rows):
                                    for mj in range(cols):
                                        cb = t.rows[mi].cells[mj]
                                        if cb is not None and cb[0] <= px <= cb[2] and cb[1] <= py <= cb[3]:
                                            master_grid[i][j] = (mi, mj)
                                            found = True
                                            break
                                    if found:
                                        break
                                if not found:
                                    # Last resort fallback
                                    if i > 0 and master_grid[i-1][j] is not None:
                                        master_grid[i][j] = master_grid[i-1][j]
                                    elif j > 0 and master_grid[i][j-1] is not None:
                                        master_grid[i][j] = master_grid[i][j-1]
                                    else:
                                        master_grid[i][j] = (i, j)
                    
                    # Build rectangular spans from master assignments
                    spans = {}
                    for i in range(rows):
                        for j in range(cols):
                            m = master_grid[i][j]
                            if m is None:
                                continue
                            if m not in spans:
                                spans[m] = [i, j]
                            else:
                                spans[m][0] = max(spans[m][0], i)
                                spans[m][1] = max(spans[m][1], j)
                    
                    for m, br in spans.items():
                        br_row, br_col = br[0], br[1]
                        if m[0] != br_row or m[1] != br_col:
                            try:
                                table_obj.cell(m[0], m[1]).merge(table_obj.cell(br_row, br_col))
                            except Exception:
                                pass
                    
                    # Format each cell
                    for i, row in enumerate(t.rows):
                        for j, cell_bbox in enumerate(row.cells):
                            if cell_bbox is None:
                                continue
                            
                            docx_cell = table_obj.cell(i, j)
                            set_cell_width(docx_cell, (cell_bbox[2] - cell_bbox[0]) * PDF_PT_TO_INCHES)
                            
                            # Detect actual alignment from PDF text positions
                            cell_words_for_align = get_words_in_bbox(words, cell_bbox)
                            detected_h, detected_v = detect_cell_alignments(cell_bbox, cell_words_for_align)
                            set_cell_valign(docx_cell, detected_v)
                            
                            # Borders
                            border_kw = {}
                            for edge_name in ('top', 'bottom', 'left', 'right'):
                                cls = classify_border(cell_bbox, edge_name, all_pdf_lines)
                                if cls == 'dotted':
                                    border_kw[edge_name] = {"val": "dotted", "sz": "4", "color": "000000", "space": "0"}
                                elif cls == 'nil':
                                    border_kw[edge_name] = {"val": "nil"}
                                else:
                                    border_kw[edge_name] = {"val": "single", "sz": "4", "color": "000000", "space": "0"}
                            set_cell_border(docx_cell, **border_kw)
                            
                            # Background
                            cell_bg_hex = None
                            max_ov = 0
                            for r in all_rects:
                                color = r.get('non_stroking_color') or r.get('fill')
                                if not color:
                                    continue
                                r_bbox = (r['x0'], r['top'], r['x1'], r['bottom'])
                                ov = overlap_area(cell_bbox, r_bbox)
                                if ov > max_ov and ov > 0.3 * bbox_area(cell_bbox):
                                    max_ov = ov
                                    cell_bg_hex = rgb_to_hex(color)
                            if cell_bg_hex and cell_bg_hex != 'FFFFFF':
                                set_cell_background(docx_cell, cell_bg_hex)
                            
                            docx_cell.text = ''
                            
                            # Images inside cell
                            for vi, vimg in enumerate(valid_images):
                                if vi in image_placed:
                                    continue
                                if vimg['bbox'] and overlap_area(cell_bbox, vimg['bbox']) > 0.3 * bbox_area(vimg['bbox']):
                                    img_stream = io.BytesIO(vimg['data'])
                                    img_p = docx_cell.paragraphs[0] if docx_cell.paragraphs else docx_cell.add_paragraph()
                                    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    try:
                                        img_w = (vimg['bbox'][2] - vimg['bbox'][0]) * PDF_PT_TO_INCHES
                                        img_p.add_run().add_picture(img_stream, width=Inches(img_w))
                                    except Exception:
                                        pass
                                    image_placed.add(vi)
                            
                            # Text
                            cw = get_words_in_bbox(words, cell_bbox)
                            if cw:
                                text_lines = group_words_into_lines(cw)
                                for li, line in enumerate(text_lines):
                                    if li == 0:
                                        p = docx_cell.paragraphs[0]
                                    else:
                                        p = docx_cell.add_paragraph()
                                    add_words_to_paragraph(p, line, cell_bbox, h_align=detected_h)
                    
                    doc.add_paragraph().paragraph_format.space_after = Pt(2)
                    last_bottom = e['bbox'][3] if 'bbox' in e else e['top'] + 20
                
                # --- FREE TEXT ---
                elif e['type'] == 'text':
                    p = doc.add_paragraph()
                    if space_before > 15:
                        p.paragraph_format.space_before = Pt(space_before - 5)
                    else:
                        p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(2)
                    
                    line_words = e['words']
                    if not line_words:
                        continue
                    
                    page_center = page.width / 2 if page.width else 300
                    line_center = (line_words[0]['x0'] + line_words[-1]['x1']) / 2
                    if abs(line_center - page_center) < 30:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        indent = line_words[0]['x0'] - 36 # 36 is the 0.5 inch margin
                        if indent > 0:
                            p.paragraph_format.left_indent = Pt(indent)
                    
                    add_words_to_paragraph(p, line_words, h_align=None)
                    last_bottom = e['bbox'][3]
            
            # Place any remaining images that weren't in a cell
            for vi, vimg in enumerate(valid_images):
                if vi not in image_placed:
                    img_stream = io.BytesIO(vimg['data'])
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    try:
                        p.add_run().add_picture(img_stream, width=Inches(2.0))
                    except Exception:
                        pass
                    image_placed.add(vi)
            
            if page_num < len(pdf.pages) - 1:
                doc.add_page_break()
    
    doc.save(output_path)
    return str(max(word_count, 1))
